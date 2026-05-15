"""
Markdown to Word converter for VPS-AR project documents.
Font: 仿宋 (SimFang), Size: 小四 (12pt)
"""

import os
import re
import glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 全局样式配置 ──────────────────────────────────────────────
FONT_BODY = '仿宋'
FONT_HEADING = '黑体'
FONT_CODE = 'Consolas'
PT_BODY = Pt(12)            # 小四 = 12pt
PT_H1 = Pt(22)              # 小二
PT_H2 = Pt(16)              # 三号
PT_H3 = Pt(14)              # 四号
PT_H4 = Pt(12)              # 小四
PT_TABLE = Pt(10.5)         # 五号
MARGIN = Cm(2.54)           # 标准页边距 1 inch


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_format(para, space_before=0, space_after=0, line_spacing=1.5,
                         first_line_indent=None, alignment=None):
    """Configure paragraph formatting."""
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    if alignment is not None:
        pf.alignment = alignment


def add_run_with_font(para, text, font_name=FONT_BODY, font_size=PT_BODY,
                      bold=False, italic=False, color=None):
    """Add a run with specified font settings."""
    run = para.add_run(text)
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run


def add_inline_formatted_text(para, text, font_name=FONT_BODY, font_size=PT_BODY):
    """Parse inline markdown formatting (bold, italic, code) and add runs."""
    # Pattern: **bold**, *italic*, `code`, plain text
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0

    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            if plain:
                add_run_with_font(para, plain, font_name, font_size)

        full, bold_text, italic_text, code_text = match.group(0), match.group(2), match.group(3), match.group(4)

        if bold_text:
            add_run_with_font(para, bold_text, font_name, font_size, bold=True)
        elif italic_text:
            add_run_with_font(para, italic_text, font_name, font_size, italic=True)
        elif code_text:
            add_run_with_font(para, code_text, FONT_CODE, Pt(font_size.pt - 1),
                              color=RGBColor(0xC7, 0x25, 0x4E))

        last_end = match.end()

    # Remaining plain text
    remaining = text[last_end:]
    if remaining:
        add_run_with_font(para, remaining, font_name, font_size)


def add_heading(doc, text: str, level: int):
    """Add a heading using Word built-in heading styles.
    Level mapping: ## -> Heading 1, ### -> Heading 2, #### -> Heading 3, etc.
    # (level 1) is ignored by the caller.
    """
    clean = text.lstrip('#').strip()
    if not clean:
        return

    # ## 映射为 Heading 1，### 映射为 Heading 2，以此类推
    word_level = min(level - 1, 4)
    style_name = f'Heading {word_level}'
    para = doc.add_paragraph(style=style_name)

    if word_level == 1:
        font_size, bold, alignment = PT_H1, True, WD_ALIGN_PARAGRAPH.CENTER
        space_before, space_after = 24, 12
    elif word_level == 2:
        font_size, bold, alignment = PT_H2, True, None
        space_before, space_after = 18, 8
    elif word_level == 3:
        font_size, bold, alignment = PT_H3, True, None
        space_before, space_after = 12, 6
    else:
        font_size, bold, alignment = PT_H4, True, None
        space_before, space_after = 8, 4

    set_paragraph_format(para, space_before, space_after, 1.5, alignment=alignment)
    add_run_with_font(para, clean, FONT_HEADING, font_size, bold=bold)


def add_paragraph_text(doc, text: str):
    """Add a normal paragraph."""
    text = text.strip()
    if not text:
        doc.add_paragraph()  # blank line
        return

    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=0, space_after=0, line_spacing=1.5,
                         first_line_indent=Cm(0.74))  # 两个字符缩进
    add_inline_formatted_text(para, text, FONT_BODY, PT_BODY)


def add_list_item(doc, text: str, ordered: bool = False):
    """Add a list item as a normal paragraph (no bullet marker)."""
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=0, space_after=2, line_spacing=1.5,
                         first_line_indent=Cm(0.74))
    # Clean leading markers
    text = re.sub(r'^[-*+]\s+', '', text)
    text = re.sub(r'^\d+\.\s+', '', text)
    add_inline_formatted_text(para, text, FONT_BODY, PT_BODY)


def add_table(doc, rows: list):
    """Parse markdown table rows and add a Word table."""
    if not rows:
        return

    # Parse rows
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip('|').split('|')]
        parsed.append(cells)

    # Filter out separator row (---)
    parsed = [r for r in parsed if not all(re.match(r'^[-:]+$', c) for c in r)]

    if not parsed:
        return

    num_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row_data in enumerate(parsed):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.cell(i, j)
            cell.text = ''
            para = cell.paragraphs[0]
            set_paragraph_format(para, space_before=2, space_after=2, line_spacing=1.15)
            clean_text = cell_text.strip()
            # Handle bold in table cells
            is_bold_header = (i == 0)
            if is_bold_header:
                add_run_with_font(para, clean_text, FONT_HEADING, PT_TABLE, bold=True)
                set_cell_shading(cell, "D9E2F3")
            else:
                add_inline_formatted_text(para, clean_text, FONT_BODY, PT_TABLE)

    # Add spacing after table
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, space_before=4, space_after=4, line_spacing=1.0)


def add_code_block(doc, lines: list):
    """Add a code block with monospace font and background."""
    for line in lines:
        para = doc.add_paragraph()
        set_paragraph_format(para, space_before=0, space_after=0, line_spacing=1.0)
        add_run_with_font(para, line, FONT_CODE, Pt(9), color=RGBColor(0x33, 0x33, 0x33))


def add_horizontal_rule(doc):
    """Add a horizontal rule (thin line)."""
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=6, space_after=6, line_spacing=1.0)
    # Use border-bottom on paragraph
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def convert_md_to_docx(md_path: str, docx_path: str):
    """Convert a single Markdown file to Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    # ── Default font ──
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = PT_BODY
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    style.paragraph_format.line_spacing = 1.5

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # ── Code block toggle ──
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table and table_rows:
                    add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r'^---+\s*$', stripped):
            if in_table and table_rows:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Heading ──
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            # Flush pending table
            if in_table and table_rows:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False
            level = len(heading_match.group(1))
            if level >= 2:  # 忽略 # 一级标题，从 ## 开始映射为 Heading 1
                add_heading(doc, stripped, level)
            i += 1
            continue

        # ── Table row ──
        if '|' in stripped and stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            i += 1
            continue
        else:
            # Flush table if we were in one
            if in_table and table_rows:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False

        # ── Empty line ──
        if not stripped:
            i += 1
            continue

        # ── Bullet list ──
        if re.match(r'^[-*+]\s+', stripped):
            add_list_item(doc, stripped, ordered=False)
            i += 1
            continue

        # ── Numbered list ──
        if re.match(r'^\d+\.\s+', stripped):
            add_list_item(doc, stripped, ordered=True)
            i += 1
            continue

        # ── Signature line (____) ──
        if re.match(r'^_{3,}', stripped):
            para = doc.add_paragraph()
            set_paragraph_format(para, space_before=6, space_after=6, line_spacing=1.5)
            add_run_with_font(para, stripped, FONT_BODY, PT_BODY)
            i += 1
            continue

        # ── Normal paragraph ──
        # Collect continuation lines (paragraph may span multiple markdown lines)
        para_lines = [stripped]
        while i + 1 < len(lines):
            next_line = lines[i + 1].rstrip('\n')
            next_stripped = next_line.strip()
            if (not next_stripped
                    or next_stripped.startswith('#')
                    or next_stripped.startswith('```')
                    or next_stripped.startswith('---')
                    or next_stripped.startswith('|')
                    or re.match(r'^[-*+]\s+', next_stripped)
                    or re.match(r'^\d+\.\s+', next_stripped)
                    or re.match(r'^_{3,}', next_stripped)):
                break
            para_lines.append(next_stripped)
            i += 1

        full_text = ' '.join(para_lines)
        add_paragraph_text(doc, full_text)
        i += 1

    # Flush any remaining table
    if in_table and table_rows:
        add_table(doc, table_rows)

    doc.save(docx_path)
    return docx_path


def main():
    """Convert all markdown files in docs/ to Word documents."""
    docs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs')
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'word')
    os.makedirs(output_dir, exist_ok=True)

    md_files = sorted(glob.glob(os.path.join(docs_dir, '*.md')))

    if not md_files:
        print("No markdown files found in docs/ directory.")
        return

    print(f"Found {len(md_files)} markdown files. Converting...\n")

    for md_path in md_files:
        basename = os.path.splitext(os.path.basename(md_path))[0]
        docx_path = os.path.join(output_dir, f'{basename}.docx')

        try:
            convert_md_to_docx(md_path, docx_path)
            print(f"  [OK] {basename}.md -> {basename}.docx")
        except Exception as e:
            print(f"  [FAIL] {basename}.md -> ERROR: {e}")

    print(f"\nDone! Word files saved to: {output_dir}")


if __name__ == '__main__':
    main()
